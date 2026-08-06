# -*- coding: utf-8 -*-
"""按项目现状修改 V1-PRD-低空经济生态服务平台.docx"""
import copy
import sys

sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

PATH = r'D:/w-yao/docs/需求文档/V1-PRD-低空经济生态服务平台.docx'
doc = Document(PATH)

# ---------- 工具函数 ----------
def set_para_text(p, new_text):
    """替换段落文本，保留第一个 run 的格式"""
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)

def set_multiline(p, lines):
    """替换段落为多行（用 <w:br/> 换行），保留第一个 run 的 rPr"""
    runs = p.runs
    template = None
    if runs:
        rpr = runs[0]._element.find(qn('w:rPr'))
        if rpr is not None:
            template = copy.deepcopy(rpr)
    for r in list(p._element.findall(qn('w:r'))):
        p._element.remove(r)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        if template is not None:
            r._element.insert(0, copy.deepcopy(template))
        if i < len(lines) - 1:
            r.add_break()

def set_cell(cell, new_text):
    """替换单元格文本（保留首段首 run 格式），删除多余段落"""
    p = cell.paragraphs[0]
    set_para_text(p, new_text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

def find_para(doc, keyword, start=0):
    """在顶层段落中按关键词查找段落对象"""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if keyword in (p.text or ''):
            return i, p
    return None, None

changed = []

def report(desc, ok):
    changed.append((desc, ok))
    print(('OK ' if ok else 'FAIL ') + desc)

# ---------- 1. 封面信息表 ----------
t = doc.tables[0]
rows = {r.cells[0].text.strip(): r for r in t.rows}
for k in rows:
    v = rows[k].cells[1].text.strip()
    print(f'  封面行[{k}] = {v}')

if '文档状态' in rows:
    set_cell(rows['文档状态'].cells[1], '开发中（V1 核心已落地）')
    report('封面: 文档状态', True)
if 'V1核心范围' in rows:
    set_cell(rows['V1核心范围'].cells[1],
             '7大业务系统 30+ 子模块（会员生态/供需对接/产学研/合规政策/人才教育/活动品牌/应急协同）+ 通用支撑')
    report('封面: V1核心范围', True)
if '技术架构建议' in rows:
    set_cell(rows['技术架构建议'].cells[1],
             '后端：Go 1.25（net/http + pgx）；数据库：PostgreSQL 16（生产）/内存存储（开发）；前端：Vue3 + Element Plus 管理后台；移动端：uni-app 微信小程序（91 页，已上线）；部署：Docker + docker-compose')
    report('封面: 技术架构', True)

# ---------- 2. 产品定位 ----------
i, p = find_para(doc, '一句话定位')
if p is not None:
    set_para_text(p, '一句话定位：面向重庆无人机产业协会180+会员单位的低空经济产业生态服务平台，已落地「会员生态、供需对接、产学研协同、合规政策、人才教育、活动品牌、应急协同」7大业务系统30+子模块，覆盖飞手接单、培训服务、供需撮合、应急调度等核心闭环，并提供认证授权、社区内容、二手交易、用工派遣、合同签约、资金托管等通用能力。')
    report('产品定位', True)
else:
    report('产品定位', False)

# ---------- 3. 功能架构总览（树形图） ----------
i, p = find_para(doc, '低空经济生态服务平台 V1')
if p is not None:
    lines = [
        '低空经济生态服务平台（7大业务系统 + 通用支撑）',
        '│',
        '├── ① 会员生态资源管控',
        '│   ├── 1. 用户与权限管理（4级RBAC + 协会7级细分）',
        '│   ├── 2. 会员单位管理（企业档案/审核/台账）',
        '│   ├── 3. 专家智库',
        '│   ├── 4. 产业资源台账（场地/试飞/中试）',
        '│   └── 5. 人才资源库',
        '│',
        '├── ② 产业供需智能对接',
        '│   ├── 6. 作业需求大厅（发布/智能匹配/竞标）',
        '│   ├── 7. 供给能力展示（整机/零部件/服务/场地）',
        '│   ├── 8. 接单与派单流程',
        '│   └── 9. 订单管理与评价',
        '│',
        '├── ③ 产学研协同创新',
        '│   ├── 10. 技术成果库',
        '│   ├── 11. 研发难题广场',
        '│   ├── 12. 课题联合攻关',
        '│   ├── 13. 中试基地测试预约',
        '│   └── 14. 成果转化追踪',
        '│',
        '├── ④ 合规政策服务',
        '│   ├── 15. 资讯与政策中心',
        '│   ├── 16. 合规知识库',
        '│   ├── 17. 团体标准管理',
        '│   ├── 18. 项目申报服务',
        '│   └── 19. 企业案例库',
        '│',
        '├── ⑤ 人才教育与产教融合',
        '│   ├── 20. 培训课程管理（课程超市）',
        '│   ├── 21. 培训报名与排班',
        '│   ├── 22. 考证管理（证书/执照）',
        '│   ├── 23. 赛事管理',
        '│   ├── 24. 招聘求职',
        '│   └── 25. 院校展示与校企共建',
        '│',
        '├── ⑥ 活动与品牌服务',
        '│   ├── 26. 活动管理',
        '│   ├── 27. 会员品牌展示',
        '│   ├── 28. 展会排期',
        '│   └── 29. 行业报告发布',
        '│',
        '├── ⑦ 低空应急资源协同',
        '│   ├── 30. 应急资源建档',
        '│   ├── 31. 一键调度',
        '│   ├── 32. 救援案例库',
        '│   ├── 33. 部门对接',
        '│   └── 34. 联合演练',
        '│',
        '└── 【通用支撑】',
        '    ├── 35. 认证授权（微信登录/RBAC/刷新令牌）',
        '    ├── 36. 消息通知（站内消息/已读未读）',
        '    ├── 37. 社区内容（发帖/评论/举报）',
        '    ├── 38. 二手交易（商品发布/收藏）',
        '    ├── 39. 用工派遣（用工订单/报价）',
        '    ├── 40. 合同签约（模板/签章回调/作废）',
        '    ├── 41. 资金托管（充值/冻结/释放/退款）',
        '    ├── 42. 文件上传（文件存储/图片优化）',
        '    └── 43. 数据统计看板（管理后台/CSV导出）',
    ]
    set_multiline(p, lines)
    report('功能架构总览树', True)
else:
    report('功能架构总览树', False)

# ---------- 4. 角色定义表 ----------
t = doc.tables[1]
hdr = [c.text.strip() for c in t.rows[0].cells]
print(f'  角色表头: {hdr}')
# 保留表头，将 6 行数据改为 5 行（4级RBAC + 协会7级）
new_rows = [
    ['platform_admin', '平台超级管理员', '协会运营团队，拥有全部权限'],
    ['association_admin', '协会管理员', '企业审核、内容管理、数据看板'],
    ['enterprise', '企业用户', '会员单位员工，可发布需求/招聘/合同/竞标报价'],
    ['individual', '个人用户', '可接单/求职/二手交易/报名培训'],
    ['协会内部7级', '会长/副会长/秘书长/部门负责人/普通会员/合作院校/访客', '会员资源管控分级浏览权限'],
]
cur = len(t.rows) - 1
for idx, row in enumerate(new_rows):
    if idx < cur:
        tr = t.rows[idx + 1]
        for j, v in enumerate(row):
            set_cell(tr.cells[j], v)
    else:
        tr = t.add_row()
        for j, v in enumerate(row):
            set_cell(tr.cells[j], v)
# 删除多余行
while len(t.rows) - 1 > len(new_rows):
    tr = t.rows[-1]._tr
    t._tbl.remove(tr)
report('角色定义表', True)

# ---------- 5. 用户注册登录（微信已上线） ----------
_, p = find_para(doc, '支持微信扫码登录')
if p is not None:
    set_para_text(p, '支持微信静默登录（小程序 code2Session，已上线，为登录主流程）；支持手机号+密码注册/登录（bcrypt 校验）')
    report('微信登录已上线', True)
else:
    report('微信登录已上线', False)

_, p = find_para(doc, '注册时选择用户类型')
if p is not None:
    set_para_text(p, '注册时选择用户类型：企业用户 / 个人用户（个人用户可接单、求职、二手交易、报名培训）')
    report('注册用户类型', True)
else:
    report('注册用户类型', False)

# ---------- 6. Non-goals 已实现项逐条更新 ----------
nongoals = [
    ('V1不做智能匹配算法推荐', '✅ 已实现：基础版智能匹配（关键词/地域/赛道推荐，/api/v1/match、/api/v1/recommendations）'),
    ('V1不做需求竞价机制', '✅ 已实现：竞标报价（demand_bids 表 + 竞标流程）'),
    ('V1不做线上支付和资金担保', '✅ 已实现：资金托管（escrow 充值/冻结/释放/退款 + 交易明细）'),
    ('V1不做合同电子签署', '✅ 已实现：电子合同（contract-templates 模板 / contracts 创建 / 签章回调 / 作废）'),
    ('V1不做订单内即时聊天', '✅ 已实现：站内消息（/api/v1/messages 列表/未读数/已读）'),
    ('V1不做课程评价', '✅ 已实现：课程评价（/api/v1/reviews/courses）'),
    ('V1不做线上支付（线下收款，系统记录状态）', '✅ 已实现：培训线上支付报名（/api/v1/training-courses/{id}/pay-and-enroll）'),
    ('V1不做数据导出', '✅ 已实现：数据导出（/api/v1/admin/export/demands、/enterprises 等 CSV）'),
    ('V1不做政策申报在线提交', '✅ 已实现：项目申报（/api/v1/project-applications 提交 + 管理后台审核）'),
]
for old, new in nongoals:
    _, p = find_para(doc, old)
    if p is not None:
        set_para_text(p, new)
        report(f'NG: {old[:12]}...', True)
    else:
        report(f'NG: {old[:12]}...', False)

# ---------- 7. 非功能需求 ----------
_, p = find_para(doc, 'API接口鉴权')
if p is not None:
    set_para_text(p, 'API接口鉴权（HMAC-SHA256 Bearer Token，Access 15分钟 / Refresh 7天轮转）')
    report('NFR: Token鉴权', True)
else:
    report('NFR: Token鉴权', False)

_, p = find_para(doc, '移动端：V1.1通过微信小程序覆盖')
if p is not None:
    set_para_text(p, '移动端：uni-app 微信小程序已上线（91 页面，Vant Weapp + 自研 u- 组件库）')
    report('NFR: 移动端', True)
else:
    report('NFR: 移动端', False)

# ---------- 8. 5.1 V1.0 范围表 ----------
# 找到范围表：表头 [模块, 功能, 优先级]
scope_tbl = None
for tb in doc.tables:
    hdrs = [c.text.strip() for c in tb.rows[0].cells]
    if hdrs == ['模块', '功能', '优先级']:
        scope_tbl = tb
        break
if scope_tbl is not None:
    new_scope = [
        ['①会员生态资源管控', '企业档案/专家智库/资源台账/人才库/协会7级权限', 'P0'],
        ['②产业供需智能对接', '需求大厅/供给展示/智能匹配/竞标报价/资源池', 'P0'],
        ['③产学研协同创新', '成果库/难题广场/课题攻关/测试预约/成果转化', 'P0'],
        ['④合规政策服务', '政策资讯/合规知识库/团体标准/项目申报/案例库', 'P0'],
        ['⑤人才教育与产教融合', '培训课程/报名排班/考证/赛事/招聘/院校共建', 'P0'],
        ['⑥活动与品牌服务', '活动管理/品牌展示/展会排期/行业报告', 'P0'],
        ['⑦低空应急资源协同', '应急资源/一键调度/救援案例/部门对接/联合演练', 'P0'],
        ['通用支撑', '认证/社区/二手/用工/合同/资金托管/消息/文件/看板', 'P0'],
    ]
    cur = len(scope_tbl.rows) - 1
    for idx, row in enumerate(new_scope):
        if idx < cur:
            tr = scope_tbl.rows[idx + 1]
            for j, v in enumerate(row):
                set_cell(tr.cells[j], v)
        else:
            tr = scope_tbl.add_row()
            for j, v in enumerate(row):
                set_cell(tr.cells[j], v)
    while len(scope_tbl.rows) - 1 > len(new_scope):
        scope_tbl._tbl.remove(scope_tbl.rows[-1]._tr)
    report('5.1 范围表', True)
else:
    report('5.1 范围表', False)

# ---------- 9. 5.2 V1.1 规划（多数已落地） ----------
_, p = find_para(doc, '微信小程序端')
if p is not None:
    set_para_text(p, '✅ 微信小程序端（uni-app，91 页，已上线）')
    report('V1.1: 小程序', True)
else:
    report('V1.1: 小程序', False)

_, p = find_para(doc, '订单内即时聊天')
if p is not None:
    set_para_text(p, '✅ 订单内即时聊天（站内消息已实现，聊天会话待完善）')
    report('V1.1: 聊天', True)
else:
    report('V1.1: 聊天', False)

_, p = find_para(doc, '需求地图可视化')
if p is not None:
    set_para_text(p, '⏳ 需求地图可视化（待开发）')
    report('V1.1: 地图', True)
else:
    report('V1.1: 地图', False)

_, p = find_para(doc, '课程在线评价')
if p is not None:
    set_para_text(p, '✅ 课程在线评价（已实现 /api/v1/reviews/courses）')
    report('V1.1: 课程评价', True)
else:
    report('V1.1: 课程评价', False)

_, p = find_para(doc, '数据导出功能')
if p is not None:
    set_para_text(p, '✅ 数据导出功能（已实现 /api/v1/admin/export/*）')
    report('V1.1: 数据导出', True)
else:
    report('V1.1: 数据导出', False)

_, p = find_para(doc, '人脸活体认证')
if p is not None:
    set_para_text(p, '⏳ 人脸活体认证（待接入第三方）')
    report('V1.1: 人脸活体', True)
else:
    report('V1.1: 人脸活体', False)

_, p = find_para(doc, '评价回复功能')
if p is not None:
    set_para_text(p, '⏳ 评价回复功能（待开发）')
    report('V1.1: 评价回复', True)
else:
    report('V1.1: 评价回复', False)

# ---------- 10. 5.3 V2.0 规划（多数已提前落地） ----------
v2 = [
    ('产学研协同创新', '✅ 产学研协同创新（技术成果库/中试基地预约/课题联合攻关）已实现'),
    ('供需智能匹配算法', '✅ 供需智能匹配算法（基础版已实现，算法优化待迭代）'),
    ('低空应急资源协同调度', '✅ 低空应急资源协同调度（应急资源/一键调度/联合演练）已实现'),
    ('行业舆情监控', '⏳ 行业舆情监控（评审已移除：与市场化产业服务定位不符，不纳入平台）'),
    ('线上支付与资金担保', '✅ 线上支付与资金担保（escrow 资金托管已实现）'),
    ('团体标准管理', '✅ 团体标准管理（compliance-standards 已实现）'),
    ('行业白皮书发布', '✅ 行业白皮书发布（industry-reports 已实现）'),
]
for old, new in v2:
    _, p = find_para(doc, old)
    if p is not None:
        set_para_text(p, new)
        report(f'V2: {old[:10]}...', True)
    else:
        report(f'V2: {old[:10]}...', False)

# ---------- 11. 业务规则 BR-9 ----------
# 找到业务规则表（表头 [编号, 规则, 说明]）
br_tbl = None
for tb in doc.tables:
    hdrs = [c.text.strip() for c in tb.rows[0].cells]
    if hdrs == ['编号', '规则', '说明']:
        br_tbl = tb
        break
if br_tbl is not None:
    for r in br_tbl.rows[1:]:
        cells = r.cells
        if cells[0].text.strip() == 'BR-9':
            set_cell(cells[1], '培训/交易付费经平台资金托管账户')
            set_cell(cells[2], '充值/冻结/释放/退款全流程托管，保障双方权益')
            report('BR-9 资金托管', True)
            break
else:
    report('BR-9 资金托管', False)

# ---------- 12. 开放问题表 ----------
oq_tbl = None
for tb in doc.tables:
    hdrs = [c.text.strip() for c in tb.rows[0].cells]
    if hdrs == ['编号', '问题', '当前处理', '待确认']:
        oq_tbl = tb
        break
if oq_tbl is not None:
    oq_updates = {
        'Q-1': ['飞手是否必须关联会员企业？', '已确认：个人用户（individual）可独立注册接单', '无需确认'],
        'Q-2': ['培训课程是否仅限协会授权机构发布？', '已确认：审核通过的会员企业/机构可发布（instructor 认证）', '无需确认'],
        'Q-3': ['需求发布是否仅限会员企业？', '已确认：企业用户可发布，管理后台审核', '无需确认'],
        'Q-4': ['短信服务商选择', '当前仅站内消息，短信通知待接入', '需确认短信签名与模板'],
        'Q-5': ['文件存储方案', '本地 uploads/ 目录（Docker 卷挂载）', '云存储（OSS/COS）待接入'],
    }
    for r in oq_tbl.rows[1:]:
        key = r.cells[0].text.strip()
        if key in oq_updates:
            vals = oq_updates[key]
            for j in range(1, 4):
                set_cell(r.cells[j], vals[j - 1])
    report('开放问题表', True)
else:
    report('开放问题表', False)

doc.save(PATH)
print()
print('=' * 40)
ok_n = sum(1 for _, ok in changed if ok)
print(f'完成: {ok_n}/{len(changed)} 项修改成功')
for desc, ok in changed:
    if not ok:
        print(f'  FAIL: {desc}')
